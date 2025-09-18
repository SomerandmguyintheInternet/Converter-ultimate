# app.py
# Universal File Utility (UFU) - Enterprise Edition v2.3 (Merged & Enhanced)
#
# This version introduces interactive table processing capabilities, including
# a live column selector and a strict table extraction mode.
#
# Features:
# - NEW: Interactive "Preview & Select Columns" tool for PDF and Word conversions.
# - NEW: "Strict Table Extraction" option for cleaner, table-only output.
# - Dual UI: Simple Mode for quick tasks & Advanced Mode for chained workflows.
# - New Conversions: Word-to-Excel, advanced Excel-to-Excel (sheet/file merging), Text file combination.
# - Expanded AI Support: Anthropic (Claude) and Grok (via Groq) support.
# - Standalone AI Tool: "AI Document Query" window to analyze and summarize any file on demand.
#
# --- REQUIRED DEPENDENCIES ---
# You MUST install these for the application to function correctly.
# Base: PyQt6, pandas, python-docx, PyPDF2, reportlab, openpyxl, "camelot-py[cv]", pyarrow
# Advanced: PyQt6-sip, qtconsole, ipykernel, psutil, pyqtgraph, qtawesome, cryptography
# AI & Others: google-generativeai, openai, requests, anthropic, groq, torch (optional, for GPU)

# =====================================================================================
# SECTION 1: IMPORTS & GLOBAL CONFIGURATION
# =====================================================================================

import sys
import os
import argparse
import threading
import re
import tempfile
import shutil
import json
import socket
import hashlib
import pickle
import queue
import sqlite3
import csv
import time
from datetime import datetime, timezone
from concurrent.futures import ThreadPoolExecutor
from typing import List, Tuple, Optional, Set, Dict, Any, IO, Type
from io import StringIO, BytesIO

# --- Core Application & UI Imports (PyQt6) ---
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QListWidget, QListWidgetItem, QLineEdit,
    QProgressBar, QTextEdit, QMessageBox, QDialog, QComboBox, QFormLayout,
    QSplitter, QGroupBox, QCheckBox, QSpinBox, QRadioButton, QMenu, QStackedWidget,
    QSizePolicy, QTabWidget, QInputDialog, QSystemTrayIcon, QDockWidget,
    QTableWidget, QTableWidgetItem, QHeaderView, QAbstractItemView
)
from PyQt6.QtCore import Qt, pyqtSignal, QObject, QSettings, QCoreApplication, QPoint, QThread, QMetaObject, Q_ARG
from PyQt6.QtGui import QAction, QIcon, QColor

# --- Safe Dependency Imports ---
try:
    import qtawesome as qta
    QTAWESOME_AVAILABLE = True
except ImportError:
    QTAWESOME_AVAILABLE = False

try:
    import torch
    TORCH_AVAILABLE = True
except ImportError:
    TORCH_AVAILABLE = False

try:
    from cryptography.exceptions import InvalidTag
    CRYPTOGRAPHY_AVAILABLE = True
except ImportError:
    CRYPTOGRAPHY_AVAILABLE = False

try:
    import pandas as pd
    from docx import Document
    from PyPDF2 import PdfReader
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    import camelot
    CONVERSION_LIBS_AVAILABLE = True
except ImportError:
    CONVERSION_LIBS_AVAILABLE = False

try:
    import google.generativeai as genai
    import openai
    import requests
    import anthropic
    import groq
    AI_LIBS_AVAILABLE = True
except ImportError:
    AI_LIBS_AVAILABLE = False

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

try:
    import pyqtgraph as pg
    PYQTGRAPH_AVAILABLE = True
except ImportError:
    PYQTGRAPH_AVAILABLE = False

try:
    from qtconsole.rich_jupyter_widget import RichJupyterWidget
    from qtconsole.inprocess import QtInProcessKernelManager
    QTCONSOLE_AVAILABLE = True
except ImportError:
    QTCONSOLE_AVAILABLE = False

# --- Global Configuration ---
APP_TITLE = "Universal File Utility (UFU) - Enterprise Edition"
SETTINGS_COMPANY = "UFUProject"
SETTINGS_APP = "UniversalFileUtility"
DB_FILE = os.path.join(os.path.expanduser("~"), ".ufu", "ufu_history.db")
DEFAULT_OUTDIR = os.path.join(os.path.expanduser("~"), "Documents", "UFU_Output")
os.makedirs(os.path.dirname(DB_FILE), exist_ok=True)


# =====================================================================================
# SECTION 2: CORE LOGIC & DATA MODELS (FROM ADVANCED SCRIPT)
# =====================================================================================

class Job:
    """Represents a single unit of work for the processing queue."""
    def __init__(self, job_id: str, source_files: List[str], workflow_steps: List['WorkflowStep'], output_dir: str):
        self.job_id = job_id
        self.source_files = source_files
        self.workflow_steps = workflow_steps
        self.output_dir = output_dir
        self.status = "queued"
        self.current_step_index = 0
        self.error_message: Optional[str] = None

    def get_display_name(self) -> str:
        if not self.source_files: return f"Job {self.job_id[:8]}"
        base_name = os.path.basename(self.source_files[0])
        return f"{base_name} (+{len(self.source_files) - 1})" if len(self.source_files) > 1 else base_name

class DatabaseManager:
    """Handles all SQLite database interactions for history and auditing."""
    def __init__(self, db_path: str):
        self.db_path = db_path
        self._create_tables_if_not_exist()

    def _get_connection(self) -> sqlite3.Connection:
        return sqlite3.connect(self.db_path)

    def _create_tables_if_not_exist(self):
        with self._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS JobHistory (
                    job_id TEXT PRIMARY KEY, job_name TEXT, status TEXT, start_time TEXT,
                    end_time TEXT, source_files_json TEXT, workflow_json TEXT,
                    output_dir TEXT, error_message TEXT
                )
            ''')
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS AuditLog (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, timestamp TEXT, event_type TEXT, details TEXT
                )
            ''')

    def log_audit_event(self, event_type: str, details: str = ""):
        with self._get_connection() as conn:
            conn.execute("INSERT INTO AuditLog (timestamp, event_type, details) VALUES (?, ?, ?)",
                         (datetime.now(timezone.utc).isoformat(), event_type, details))

    def log_job_start(self, job: Job):
        with self._get_connection() as conn:
            conn.execute('''
                INSERT INTO JobHistory (job_id, job_name, status, start_time, source_files_json, workflow_json, output_dir)
                VALUES (?, ?, 'running', ?, ?, ?, ?)
            ''', (job.job_id, job.get_display_name(), datetime.now(timezone.utc).isoformat(),
                  json.dumps(job.source_files), json.dumps([step.to_dict() for step in job.workflow_steps]), job.output_dir))

    def log_job_completion(self, job: Job):
        with self._get_connection() as conn:
            conn.execute('''
                UPDATE JobHistory SET status = ?, end_time = ?, error_message = ? WHERE job_id = ?
            ''', (job.status, datetime.now(timezone.utc).isoformat(), job.error_message, job.job_id))

    def get_job_history(self, limit: int = 100) -> List[Dict[str, Any]]:
        with self._get_connection() as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM JobHistory ORDER BY start_time DESC LIMIT ?", (limit,))
            return [dict(row) for row in cursor.fetchall()]

class AIProcessor:
    """Handles all interactions with an external LLM API (Gemini, OpenAI, Claude, Grok)."""
    def __init__(self, signals: 'WorkerSignals'):
        self.signals = signals; self.model = None; self.client = None; self.api_type = None; self.chat_session = None

    def is_ready(self) -> bool:
        """Check if the processor is initialized and ready to use."""
        return self.client is not None or self.model is not None

    def initialize(self, api_type: str, api_key: str) -> bool:
        self.api_type = api_type
        try:
            if not AI_LIBS_AVAILABLE: raise ImportError("Required AI libraries are not installed.")
            if api_type == "Google Gemini":
                genai.configure(api_key=api_key); self.model = genai.GenerativeModel('gemini-1.5-flash')
                next(genai.list_models()) # Validation call
            elif api_type == "OpenAI":
                self.client = openai.OpenAI(api_key=api_key); self.client.models.list() # Validation call
            elif api_type == "Anthropic Claude":
                self.client = anthropic.Anthropic(api_key=api_key); self.client.count_tokens(text="validation") # Validation call
            elif api_type == "Grok (via Groq)":
                self.client = groq.Groq(api_key=api_key); self.client.models.list() # Validation call
            else: raise ValueError("Unsupported AI API type")
            self.signals.log.emit(f"✅ AI Processor initialized successfully for {api_type}.")
            return True
        except Exception as e:
            self.model = self.client = None
            self.signals.log.emit(f"🔴 <font color='red'>AI Processor failed to initialize:</font> {e}")
            return False

    def execute_single_prompt(self, prompt: str) -> str:
        if not self.is_ready(): return "Error: AI Processor not initialized."
        try:
            if self.api_type == "Google Gemini":
                return self.model.generate_content(prompt).text
            elif self.api_type == "OpenAI":
                completion = self.client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}])
                return completion.choices[0].message.content
            elif self.api_type == "Anthropic Claude":
                message = self.client.messages.create(model="claude-3-haiku-20240307", max_tokens=4096, messages=[{"role": "user", "content": prompt}])
                return message.content[0].text
            elif self.api_type == "Grok (via Groq)":
                completion = self.client.chat.completions.create(model="llama3-8b-8192", messages=[{"role": "user", "content": prompt}])
                return completion.choices[0].message.content
        except Exception as e: return f"Error during AI execution: {e}"
        return "An unknown error occurred."

# =====================================================================================
# SECTION 3: CORE LOGIC & PROCESSORS (FROM ORIGINAL APP.PY)
# =====================================================================================

class FunModeProcessor:
    TAG_HEADER = b'UFCX-TAG-V2    \x00\x00'; HEADER_VERSION = 0x02; XOR_KEY = b'a-better-secret-key-for-fun-mode'
    def __init__(self, device: str = "auto"):
        if not TORCH_AVAILABLE: self.device = "cpu"
        elif device == "auto": self.device = "cuda:0" if torch.cuda.is_available() else "cpu"
        else: self.device = device
        if self.device.startswith("cuda") and not torch.cuda.is_available(): self.device = "cpu"

    def _xor_cipher_stream(self, fin: IO[bytes], fout: IO[bytes], chunk_size: int):
        use_gpu = self.device.startswith("cuda")
        if use_gpu: key_tensor = torch.tensor(list(self.XOR_KEY), dtype=torch.uint8, device=self.device)
        while chunk := fin.read(chunk_size):
            if use_gpu:
                try:
                    data_tensor = torch.tensor(list(chunk), dtype=torch.uint8, device=self.device)
                    key_expanded = key_tensor.repeat(len(data_tensor) // len(self.XOR_KEY) + 1)[:len(data_tensor)]
                    fout.write(bytes((data_tensor ^ key_expanded).cpu().numpy().tolist()))
                except Exception: fout.write(bytes(b ^ self.XOR_KEY[i % len(self.XOR_KEY)] for i, b in enumerate(chunk)))
            else: fout.write(bytes(b ^ self.XOR_KEY[i % len(self.XOR_KEY)] for i, b in enumerate(chunk)))

    def encrypt_file(self, input_path: str, output_path: str, chunk_size: int) -> bool:
        try:
            metadata = {"orig_name": os.path.basename(input_path), "created_utc": datetime.now(timezone.utc).isoformat()}
            metadata_bytes = json.dumps(metadata).encode('utf-8')
            with open(output_path, 'wb') as fout, open(input_path, 'rb') as fin:
                fout.write(self.TAG_HEADER); fout.write(self.HEADER_VERSION.to_bytes(1, 'big'))
                fout.write(len(metadata_bytes).to_bytes(2, 'big')); fout.write(metadata_bytes)
                self._xor_cipher_stream(fin, fout, chunk_size)
            return True
        except Exception: return False

    def decrypt_file_stream(self, input_path: str, output_path: str, chunk_size: int) -> Optional[dict]:
        try:
            with open(input_path, 'rb') as fin:
                if fin.read(16) != self.TAG_HEADER or int.from_bytes(fin.read(1), 'big') != self.HEADER_VERSION: return None
                metadata_len = int.from_bytes(fin.read(2), 'big')
                metadata = json.loads(fin.read(metadata_len).decode('utf-8'))
                with open(output_path, 'wb') as fout: self._xor_cipher_stream(fin, fout, chunk_size)
                return metadata
        except Exception: return None

class EmbedProcessor:
    EMBED_TAG_FOOTER = b'UFC-EMBED-V1\x00\x00\x00\x00'; FOOTER_LEN = len(EMBED_TAG_FOOTER)
    def embed_file(self, host_path: str, payload_path: str, output_path: str) -> bool:
        try:
            with open(payload_path, 'rb') as f: payload_content = f.read()
            metadata = {"payload_name": os.path.basename(payload_path), "payload_len": len(payload_content)}
            metadata_bytes = json.dumps(metadata).encode('utf-8')
            shutil.copyfile(host_path, output_path)
            with open(output_path, 'ab') as f:
                f.write(payload_content); f.write(metadata_bytes)
                f.write(len(metadata_bytes).to_bytes(4, 'big')); f.write(self.EMBED_TAG_FOOTER)
            return True
        except Exception: return False

    def extract_embedded_file_info(self, input_path: str) -> Optional[Tuple[dict, bytes]]:
        try:
            with open(input_path, 'rb') as f:
                f.seek(-self.FOOTER_LEN, os.SEEK_END)
                if f.read(self.FOOTER_LEN) != self.EMBED_TAG_FOOTER: return None
                f.seek(-(self.FOOTER_LEN + 4), os.SEEK_END); metadata_len = int.from_bytes(f.read(4), 'big')
                f.seek(-(self.FOOTER_LEN + 4 + metadata_len), os.SEEK_END); metadata = json.loads(f.read(metadata_len).decode('utf-8'))
                payload_len = metadata.get("payload_len")
                if payload_len is None: return None
                f.seek(-(self.FOOTER_LEN + 4 + metadata_len + payload_len), os.SEEK_END)
                return metadata, f.read(payload_len)
        except Exception: return None

def clean_bank_statement_df(df: pd.DataFrame) -> pd.DataFrame:
    # This complex heuristic function is preserved as-is from the original script
    header_keywords = {'date': ['date'], 'description': ['description', 'remarks'], 'money': ['debit', 'credit', 'balance', 'amount']}
    header_row_index = -1
    for i, row in df.iterrows():
        row_str = ' '.join(map(str, row.values)).lower()
        if any(kw in row_str for kw in header_keywords['date']) and any(kw in row_str for kw in header_keywords['money']):
            header_row_index = i; break
    if header_row_index == -1:
        # If no header found, assume first row is header and try to clean
        df.columns = df.iloc[0]
        df = df[1:]
    else:
        header = df.iloc[header_row_index].astype(str).str.lower().str.strip()
        df_cleaned = df.iloc[header_row_index + 1:].copy(); df_cleaned.columns = header
        df = df_cleaned

    column_map = {'txn date': 'date', 'transaction date': 'date', 'narration': 'description', 'remarks': 'description', 'withdrawal': 'debit', 'deposit': 'credit'}
    df.rename(columns=lambda c: column_map.get(str(c).lower().strip(), str(c).lower().strip()), inplace=True)

    if 'date' not in df.columns: return pd.DataFrame() # Date is essential
    
    money_cols = [c for c in ['debit', 'credit', 'balance', 'amount'] if c in df.columns]
    for col in money_cols:
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(r'[^\d.-]', '', regex=True), errors='coerce').fillna(0)
    
    df['date'] = pd.to_datetime(df['date'], errors='coerce', dayfirst=True).dt.strftime('%Y-%m-%d')
    output_rename_map = {'date': 'Date', 'description': 'Remarks', 'debit': 'Withdrawals', 'credit': 'Deposits', 'balance': 'Balance', 'amount':'Amount'}
    df.rename(columns=output_rename_map, inplace=True)
    
    final_cols = [col for col in output_rename_map.values() if col in df.columns]
    return df.dropna(subset=['Date'])[final_cols] if final_cols else pd.DataFrame()


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            reader = PdfReader(path)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        if ext in (".xls", ".xlsx"):
            return "\n".join([df.to_csv(index=False) for df in pd.read_excel(path, sheet_name=None).values()])
        if ext == ".docx": return "\n".join(p.text for p in Document(path).paragraphs)
        if ext in (".txt", ".csv", ".md", ".json", ".xml", ".html"):
            return open(path, "r", encoding="utf-8", errors="ignore").read()
        return f"[ERROR: File type '{ext}' is not supported for text extraction]"
    except Exception as e: return f"[ERROR reading {path}: {e}]"

# =====================================================================================
# SECTION 4: WORKFLOW STEPS (NEW MODULAR ARCHITECTURE)
# =====================================================================================

class WorkflowStep:
    name = "Abstract Step"; description = "Placeholder step."
    def __init__(self): self.config = self.get_default_config()
    def get_default_config(self) -> Dict[str, Any]: return {}
    def get_config_widget(self) -> QWidget:
        widget = QWidget(); QVBoxLayout(widget).addWidget(QLabel(f"No specific options for '{self.name}'."))
        return widget
    def update_config_from_widget(self, widget: QWidget): pass
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        raise NotImplementedError
    def to_dict(self) -> Dict[str, Any]: return {'name': self.name, 'config': self.config}
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'WorkflowStep':
        step = cls(); step.config.update(data.get('config', {}))
        return step

# --- New & Enhanced Conversion Steps ---

class PdfToExcelStep(WorkflowStep):
    name = "Convert: PDF to Excel"; description = "Extracts tables from PDFs (optimized for statements) into one Excel file."
    # This step is a batch processor, handled specially by the worker.
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        return input_path # Pass through

class WordToExcelStep(WorkflowStep):
    name = "Convert: Word to Excel"; description = "Extracts all tables from a .docx file into sheets of an Excel file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        if not input_path.lower().endswith(".docx"): raise ValueError("This step only accepts .docx files.")
        doc = Document(input_path)
        if not doc.tables: raise RuntimeError("No tables found in the document.")
        
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}_tables.xlsx")
        
        selected_columns = job_settings.get('selected_columns')

        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            for i, table in enumerate(doc.tables):
                data = [[cell.text for cell in row.cells] for row in table.rows]
                if not data: continue

                df = pd.DataFrame(data[1:], columns=data[0]) # Assume first row is header
                if selected_columns:
                    # Filter for columns that actually exist in this specific table
                    existing_cols_in_order = [col for col in selected_columns if col in df.columns]
                    if existing_cols_in_order:
                        df = df[existing_cols_in_order]

                df.to_excel(writer, sheet_name=f'Table_{i+1}', index=False)
        signals.log.emit(f"    - Extracted {len(doc.tables)} tables from {os.path.basename(input_path)}")
        return output_path

class ExcelProcessStep(WorkflowStep):
    name = "Process: Excel File"; description = "Processes Excel files, e.g., combining all sheets into one."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        if not self.config.get('combine_sheets'): return input_path # No operation needed for single file
        xls = pd.ExcelFile(input_path)
        if len(xls.sheet_names) <= 1: return input_path # No need to combine
        all_sheets_df = pd.concat([pd.read_excel(xls, sheet_name=name) for name in xls.sheet_names], ignore_index=True)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}_sheets_combined.xlsx")
        all_sheets_df.to_excel(output_path, index=False)
        signals.log.emit(f"    - Combined {len(xls.sheet_names)} sheets from {os.path.basename(input_path)}")
        return output_path

class CombineFilesStep(WorkflowStep):
    name = "Combine Files"; description = "Combines multiple text or Excel files into a single output file."
    # This step is a batch processor, handled specially by the worker.
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        return input_path # Pass through

# --- Original Steps ---

class FunEncryptStep(WorkflowStep):
    name = "Encrypt (Fun Mode)"; description = "Disguises a file using a simple, reversible XOR cipher."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        processor = FunModeProcessor(device=job_settings.get('gpu_device', "auto"))
        fake_ext = job_settings.get('force_ext', ".fun")
        chunk_size = job_settings.get('chunk_size_mb', 64) * 1024 * 1024
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, base_name + fake_ext)
        if not processor.encrypt_file(input_path, output_path, chunk_size):
            raise RuntimeError("Fun Encryption process failed internally.")
        signals.log.emit(f"    - Fun Encrypted: {os.path.basename(input_path)} -> {os.path.basename(output_path)}")
        return output_path

class FunDecryptStep(WorkflowStep):
    name = "Decrypt (Fun Mode)"; description = "Reverses the 'Fun Mode' XOR cipher to recover the original file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        processor = FunModeProcessor(device=job_settings.get('gpu_device', "auto"))
        chunk_size = job_settings.get('chunk_size_mb', 64) * 1024 * 1024
        temp_output_path = os.path.join(output_dir, f"temp_{os.path.basename(input_path)}")
        metadata = processor.decrypt_file_stream(input_path, temp_output_path, chunk_size)
        if not metadata: raise RuntimeError("Not a valid Fun-Encrypted file or is corrupted.")
        original_name = metadata.get("orig_name", f"decrypted_{os.path.basename(input_path)}")
        final_output_path = os.path.join(output_dir, original_name)
        shutil.move(temp_output_path, final_output_path)
        signals.log.emit(f"    - Fun Decrypted: {os.path.basename(input_path)} -> {original_name}")
        return final_output_path

class EmbedFileStep(WorkflowStep):
    name = "Embed File"; description = "Hides a payload file inside a host file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        host_path = input_path
        payload_path = job_settings.get('payload_file')
        if not payload_path or not os.path.exists(payload_path): raise ValueError("Payload file for embedding not found.")
        processor = EmbedProcessor()
        output_path = os.path.join(output_dir, f"embedded_{os.path.basename(host_path)}")
        if not processor.embed_file(host_path, payload_path, output_path):
            raise RuntimeError("File embedding failed.")
        signals.log.emit(f"    - Embedded: {os.path.basename(payload_path)} into {os.path.basename(output_path)}")
        return output_path

class ExtractFileStep(WorkflowStep):
    name = "Extract File"; description = "Extracts a hidden payload file from a host file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        processor = EmbedProcessor()
        result = processor.extract_embedded_file_info(input_path)
        if not result: raise RuntimeError("No embedded payload found or file is corrupted.")
        metadata, content = result
        payload_name = metadata.get("payload_name", f"extracted_from_{os.path.basename(input_path)}")
        output_path = os.path.join(output_dir, payload_name)
        with open(output_path, 'wb') as f: f.write(content)
        signals.log.emit(f"    - Extracted: {os.path.basename(input_path)} -> {payload_name}")
        return output_path

class GeneralConvertStep(WorkflowStep):
    name = "Convert: General Text Extraction"; description = "Extracts text content from various file types into a plain text file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        text_content = extract_text_from_file(input_path)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}.txt")
        with open(output_path, 'w', encoding='utf-8') as f: f.write(text_content)
        signals.log.emit(f"    - Converted: {os.path.basename(input_path)} -> {os.path.basename(output_path)}")
        return output_path

class AISummarizeStep(WorkflowStep):
    name = "AI: Summarize Document"; description = "Uses an LLM to generate a summary of a text-based file."
    def execute(self, input_path: str, output_dir: str, signals: 'WorkerSignals', job_settings: Dict[str, Any]) -> str:
        ai_processor = job_settings.get('ai_processor')
        if not ai_processor or not ai_processor.is_ready(): raise RuntimeError("AI Processor not available for summarization.")
        try: content = open(input_path, 'r', encoding='utf-8', errors='ignore').read()
        except Exception: raise ValueError(f"Could not read text from {os.path.basename(input_path)}.")
        prompt = f"Please provide a concise, 2-paragraph summary of the following document:\n\n---\n{content[:25000]}\n---"
        signals.log.emit(f"    - Sending document to AI for summarization...")
        summary_text = ai_processor.execute_single_prompt(prompt)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}_summary.txt")
        with open(output_path, 'w', encoding='utf-8') as f: f.write(summary_text)
        signals.log.emit(f"    - AI Summary saved to {os.path.basename(output_path)}")
        return output_path

ALL_WORKFLOW_STEPS: Dict[str, Type[WorkflowStep]] = {
    "Convert: General Text Extraction": GeneralConvertStep,
    "Convert: PDF to Excel": PdfToExcelStep,
    "Convert: Word to Excel": WordToExcelStep,
    "Process: Excel File": ExcelProcessStep,
    "Combine Files": CombineFilesStep,
    "Encrypt (Fun Mode)": FunEncryptStep,
    "Decrypt (Fun Mode)": FunDecryptStep,
    "Embed File": EmbedFileStep,
    "Extract File": ExtractFileStep,
    "AI: Summarize Document": AISummarizeStep,
}

# =====================================================================================
# SECTION 5: BACKGROUND WORKERS & THREADING
# =====================================================================================

class WorkerSignals(QObject):
    log = pyqtSignal(str); job_finished = pyqtSignal(Job); all_finished = pyqtSignal()
    update_job_dashboard = pyqtSignal()

class AppWorker(QObject):
    def __init__(self, job_queue: queue.Queue, signals: WorkerSignals, db_manager: DatabaseManager, ai_processor: AIProcessor):
        super().__init__()
        self.job_queue = job_queue; self.signals = signals; self.db_manager = db_manager
        self.ai_processor = ai_processor; self.is_running = True

    def run(self):
        self.db_manager.log_audit_event("Worker.Started", "Worker thread has started.")
        while self.is_running:
            try:
                job = self.job_queue.get(timeout=1)
                self.process_job(job)
                self.job_queue.task_done()
            except queue.Empty: continue
            except Exception as e: self.signals.log.emit(f"🔴 <font color='red'>Critical worker error:</font> {e}")
        self.signals.log.emit("Worker thread has stopped."); self.signals.all_finished.emit()
        self.db_manager.log_audit_event("Worker.Stopped", "Worker thread has stopped.")

    def process_job(self, job: Job):
        self.signals.log.emit(f"🚀 Starting job: <b>{job.get_display_name()}</b>")
        job.status = "running"; self.db_manager.log_job_start(job); self.signals.update_job_dashboard.emit()
        temp_dir = tempfile.mkdtemp(prefix="ufu_job_")

        try:
            # Handle special batch processing cases
            is_batch_job = len(job.workflow_steps) == 1 and isinstance(job.workflow_steps[0], (PdfToExcelStep, CombineFilesStep))
            
            if is_batch_job and isinstance(job.workflow_steps[0], PdfToExcelStep):
                self._run_specialized_pdf_process(job)
            elif is_batch_job and isinstance(job.workflow_steps[0], CombineFilesStep):
                file_type = job.workflow_steps[0].config.get('file_type', 'Text')
                if file_type == 'Excel': self._run_combine_excel_process(job)
                else: self._run_combine_text_process(job)
            else: # General workflow processing for each file
                for i, source_file in enumerate(job.source_files):
                    self.signals.log.emit(f"  - Processing file {i+1}/{len(job.source_files)}: {os.path.basename(source_file)}")
                    current_file = source_file
                    for step_idx, step in enumerate(job.workflow_steps):
                        job.current_step_index = step_idx
                        self.signals.log.emit(f"    - Executing step {step_idx+1}/{len(job.workflow_steps)}: <b>{step.name}</b>")
                        job_settings = step.config.copy()
                        job_settings['ai_processor'] = self.ai_processor
                        output_path = step.execute(current_file, temp_dir, self.signals, job_settings)
                        if output_path != current_file and os.path.exists(output_path): current_file = output_path
                    # Move final result from temp to output dir, unless it was an in-place operation
                    if os.path.dirname(current_file) == os.path.realpath(temp_dir):
                        final_output_path = os.path.join(job.output_dir, os.path.basename(current_file))
                        shutil.move(current_file, final_output_path)
                        self.signals.log.emit(f"  - ✅ Finished {os.path.basename(source_file)}. Final output: {os.path.basename(final_output_path)}")
                    else:
                        self.signals.log.emit(f"  - ✅ Finished {os.path.basename(source_file)}. No new file created.")

            job.status = "success"
            self.signals.log.emit(f"✅ <font color='green'>Job '{job.get_display_name()}' completed successfully.</font>")
        except Exception as e:
            job.status = "failed"; job.error_message = str(e)
            step_name = job.workflow_steps[job.current_step_index].name
            self.signals.log.emit(f"🔴 <font color='red'>Job '{job.get_display_name()}' failed at step '{step_name}':</font> {e}")
        finally:
            shutil.rmtree(temp_dir)
            self.db_manager.log_job_completion(job)
            self.signals.job_finished.emit(job)
            self.signals.update_job_dashboard.emit()

    def _run_specialized_pdf_process(self, job: Job):
        pdf_files = [f for f in job.source_files if f.lower().endswith('.pdf')]
        if not pdf_files: raise RuntimeError("No PDF files provided for PDF-to-Excel task.")

        step_config = job.workflow_steps[0].config
        flavor = step_config.get('flavor', 'stream')
        selected_columns = step_config.get('selected_columns')
        self.signals.log.emit(f"  - Using Camelot flavor: '{flavor}'")
        
        all_dfs = []
        for path in pdf_files:
            try:
                tables = camelot.read_pdf(path, pages='all', flavor=flavor)
                if tables.n > 0:
                    valid_dfs = [df for df in [clean_bank_statement_df(tbl.df) for tbl in tables] if not df.empty]
                    if valid_dfs:
                        # Apply column selection if specified
                        if selected_columns:
                            processed_dfs = []
                            for df in valid_dfs:
                                existing_cols = [col for col in selected_columns if col in df.columns]
                                if existing_cols:
                                    processed_dfs.append(df[existing_cols])
                            valid_dfs = processed_dfs
                        
                        if valid_dfs:
                            file_df = pd.concat(valid_dfs, ignore_index=True)
                            file_df['source_file'] = os.path.basename(path)
                            all_dfs.append(file_df)
            except Exception as e:
                self.signals.log.emit(f"  - ERROR processing {os.path.basename(path)}: {e}")

        if not all_dfs: raise RuntimeError("No data was successfully extracted from any PDF.")
        
        final_df = pd.concat(all_dfs, ignore_index=True)
        if step_config.get('sort_by_date', True) and 'Date' in final_df.columns:
            final_df['Date'] = pd.to_datetime(final_df['Date'], errors='coerce')
            final_df.sort_values(by='Date', ascending=False, inplace=True)
        
        filename = step_config.get('combine_filename', 'combined_output')
        output_path = os.path.join(job.output_dir, f"{filename}.xlsx")
        final_df.to_excel(output_path, index=False)
        self.signals.log.emit(f"  - ✅ Combined file created: {os.path.basename(output_path)}")

    def _run_combine_text_process(self, job: Job):
        step_config = job.workflow_steps[0].config
        filename = step_config.get('combine_filename', 'combined_text_output')
        output_path = os.path.join(job.output_dir, f"{filename}.txt")
        
        with open(output_path, 'w', encoding='utf-8') as outfile:
            for i, f in enumerate(job.source_files):
                self.signals.log.emit(f"  - Combining file {i+1}/{len(job.source_files)}: {os.path.basename(f)}")
                outfile.write(f"\n\n--- CONTENT FROM: {os.path.basename(f)} ---\n\n")
                try:
                    with open(f, 'r', encoding='utf-8', errors='ignore') as infile:
                        outfile.write(infile.read())
                except Exception as e:
                    outfile.write(f"[ERROR READING FILE: {e}]")
        self.signals.log.emit(f"  - ✅ Combined text file created: {os.path.basename(output_path)}")

    def _run_combine_excel_process(self, job: Job):
        excel_files = [f for f in job.source_files if f.lower().endswith(('.xlsx', '.xls'))]
        if not excel_files: raise RuntimeError("No Excel files found for combination task.")
        
        step_config = job.workflow_steps[0].config
        combine_sheets = step_config.get('combine_sheets', True)
        all_dfs = []
        for path in excel_files:
            try:
                xls = pd.ExcelFile(path)
                if combine_sheets:
                    file_df = pd.concat([pd.read_excel(xls, sheet_name=name) for name in xls.sheet_names], ignore_index=True)
                    file_df['source_file'] = os.path.basename(path)
                    all_dfs.append(file_df)
                else: # Add each sheet as a separate dataframe
                    for name in xls.sheet_names:
                        sheet_df = pd.read_excel(xls, sheet_name=name)
                        sheet_df['source_file'] = f"{os.path.basename(path)} - {name}"
                        all_dfs.append(sheet_df)
            except Exception as e:
                self.signals.log.emit(f"  - ERROR processing {os.path.basename(path)}: {e}")

        if not all_dfs: raise RuntimeError("No data was successfully extracted from any Excel file.")
        
        final_df = pd.concat(all_dfs, ignore_index=True)
        filename = step_config.get('combine_filename', 'combined_excel_output')
        output_path = os.path.join(job.output_dir, f"{filename}.xlsx")
        final_df.to_excel(output_path, index=False)
        self.signals.log.emit(f"  - ✅ Combined Excel file created: {os.path.basename(output_path)}")

    def stop(self): self.is_running = False

# =====================================================================================
# SECTION 6: GUI WIDGETS & DIALOGS
# =====================================================================================

class DropLineEdit(QLineEdit):
    def __init__(self, parent=None): super().__init__(parent); self.setAcceptDrops(True)
    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls() and e.mimeData().urls()[0].isLocalFile(): e.acceptProposedAction()
    def dropEvent(self, e): self.setText(e.mimeData().urls()[0].toLocalFile())

class FileListWidget(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent); self.setAcceptDrops(True); self.setSelectionMode(self.SelectionMode.ExtendedSelection)
    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls(): e.acceptProposedAction()
    def dropEvent(self, e):
        for u in e.mimeData().urls():
            path = u.toLocalFile()
            if os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for name in files: self.add_file_path(os.path.join(root, name))
            else: self.add_file_path(path)
    def add_file_path(self, path: str):
        if not self.findItems(path, Qt.MatchFlag.MatchExactly): self.addItem(path)
    def get_all_paths(self) -> List[str]: return [self.item(i).text() for i in range(self.count())]

class SettingsDialog(QDialog):
    def __init__(self, current_settings: Dict[str, Any], parent=None):
        super().__init__(parent); self.setWindowTitle("Global Settings"); self.setMinimumWidth(450)
        self.settings = current_settings.copy()
        layout = QFormLayout(self)
        if TORCH_AVAILABLE:
            self.gpu_combo = QComboBox()
            self.gpu_combo.addItems(["Auto (Recommended)", "CPU only"])
            current_device = self.settings.get('gpu_device', 'auto')
            if torch.cuda.is_available():
                for i in range(torch.cuda.device_count()): self.gpu_combo.addItem(f"GPU {i}: {torch.cuda.get_device_name(i)}")
            if current_device == 'cpu': self.gpu_combo.setCurrentIndex(1)
            else: self.gpu_combo.setCurrentIndex(0)
            layout.addRow("GPU Acceleration Device:", self.gpu_combo)
        
        self.chunk_size_spin = QSpinBox()
        self.chunk_size_spin.setRange(1, 4096); self.chunk_size_spin.setSuffix(" MB")
        self.chunk_size_spin.setValue(self.settings.get('chunk_size_mb', 64))
        layout.addRow("Processing Chunk Size (Encryption):", self.chunk_size_spin)

        button_box = QHBoxLayout(); ok_button = QPushButton("Save"); ok_button.clicked.connect(self.accept)
        cancel_button = QPushButton("Cancel"); cancel_button.clicked.connect(self.reject)
        button_box.addStretch(); button_box.addWidget(ok_button); button_box.addWidget(cancel_button); layout.addRow(button_box)
    
    def get_settings(self) -> Dict[str, Any]:
        if TORCH_AVAILABLE:
            sel = self.gpu_combo.currentText()
            if "CPU" in sel: self.settings['gpu_device'] = "cpu"
            elif "GPU" in sel: self.settings['gpu_device'] = f"cuda:{sel.split(':')[0].split(' ')[1]}"
            else: self.settings['gpu_device'] = "auto"
        self.settings['chunk_size_mb'] = self.chunk_size_spin.value()
        return self.settings

class AIQueryDialog(QDialog):
    def __init__(self, ai_processor: AIProcessor, parent=None):
        super().__init__(parent)
        self.ai_processor = ai_processor
        self.setWindowTitle("AI Document Query")
        self.setMinimumSize(700, 500)
        self.thread = None
        self.worker = None

        main_layout = QVBoxLayout(self)
        
        # File selection
        file_layout = QHBoxLayout()
        self.file_path_edit = DropLineEdit()
        self.file_path_edit.setPlaceholderText("Drag & drop a file or browse...")
        self.file_path_edit.textChanged.connect(self.load_file_content)
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_for_file)
        file_layout.addWidget(QLabel("File:"))
        file_layout.addWidget(self.file_path_edit)
        file_layout.addWidget(browse_btn)
        main_layout.addLayout(file_layout)

        # Query and Response
        splitter = QSplitter(Qt.Orientation.Vertical)
        
        query_group = QGroupBox("Your Query")
        query_layout = QVBoxLayout(query_group)
        self.query_edit = QLineEdit()
        self.query_edit.setPlaceholderText("e.g., Summarize this document in three bullet points.")
        self.query_edit.returnPressed.connect(self.start_ai_query)
        ask_btn = QPushButton("Ask AI")
        ask_btn.clicked.connect(self.start_ai_query)
        query_layout.addWidget(self.query_edit)
        query_layout.addWidget(ask_btn)
        splitter.addWidget(query_group)

        response_group = QGroupBox("AI Response")
        response_layout = QVBoxLayout(response_group)
        self.response_text = QTextEdit()
        self.response_text.setReadOnly(True)
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0,0) # Indeterminate
        self.progress_bar.setVisible(False)
        response_layout.addWidget(self.response_text)
        response_layout.addWidget(self.progress_bar)
        splitter.addWidget(response_group)

        main_layout.addWidget(splitter)
        splitter.setSizes([100, 400])

    def browse_for_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select a file to query")
        if path:
            self.file_path_edit.setText(path)

    def load_file_content(self, path):
        if not os.path.exists(path): return
        # For simplicity, we just show a message. A preview is too complex.
        self.response_text.setPlaceholderText(f"Ready to ask questions about:\n{os.path.basename(path)}")

    def start_ai_query(self):
        file_path = self.file_path_edit.text()
        query = self.query_edit.text()
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Error", "Please select a valid file.")
            return
        if not query:
            QMessageBox.warning(self, "Error", "Please enter a query.")
            return
            
        self.progress_bar.setVisible(True)
        self.response_text.clear()
        
        self.worker = AIQueryWorker(self.ai_processor, file_path, query)
        self.thread = QThread()
        self.worker.moveToThread(self.thread)
        self.worker.finished.connect(self.on_ai_query_finished)
        self.thread.started.connect(self.worker.run)
        self.thread.start()

    def on_ai_query_finished(self, response):
        self.response_text.setText(response)
        self.progress_bar.setVisible(False)
        self.thread.quit()
        self.thread.wait()

class AIQueryWorker(QObject):
    finished = pyqtSignal(str)
    def __init__(self, ai_processor, file_path, query):
        super().__init__()
        self.ai_processor = ai_processor
        self.file_path = file_path
        self.query = query

    def run(self):
        content = extract_text_from_file(self.file_path)
        if content.startswith("[ERROR"):
            self.finished.emit(content)
            return

        prompt = f"""You are an AI assistant analyzing a document.
        
        Document Content:
        ---
        {content[:30000]}
        ---
        
        User's Query: "{self.query}"
        
        Please provide a clear and concise answer based on the document's content.
        """
        response = self.ai_processor.execute_single_prompt(prompt)
        self.finished.emit(response)

# =====================================================================================
# SECTION 7: SIMPLE MODE UI (NOW WITH COLUMN SELECTOR)
# =====================================================================================

class SimpleModeWidget(QWidget):
    def __init__(self, main_window: 'MainWindow', parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.fun_mode_disclaimer_shown = False
        self.custom_column_config = None # Stores list of selected column names for a job
        self.init_ui()

    def init_ui(self):
        main_layout = QHBoxLayout(self)
        splitter = QSplitter(Qt.Orientation.Horizontal); main_layout.addWidget(splitter)
        
        # --- Left Pane ---
        left_widget = QWidget(); left_layout = QVBoxLayout(left_widget)
        left_layout.addWidget(QLabel("<h3>1. Add Files for Batch Processing</h3>"))
        self.file_list = FileListWidget()
        left_layout.addWidget(self.file_list)
        file_btns = QHBoxLayout(); add_btn = QPushButton("Add Files..."); add_btn.clicked.connect(self.add_files_dialog)
        add_folder_btn = QPushButton("Add Folder..."); add_folder_btn.clicked.connect(self.add_folder_dialog)
        clear_btn = QPushButton("Clear"); clear_btn.clicked.connect(self.file_list.clear)
        file_btns.addStretch(); file_btns.addWidget(add_btn); file_btns.addWidget(add_folder_btn); file_btns.addWidget(clear_btn)
        left_layout.addLayout(file_btns); splitter.addWidget(left_widget)

        # --- Right Pane ---
        right_widget = QWidget(); right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("<h3>2. Configure and Run Operation</h3>"))

        mode_group = QGroupBox("Select Operation Mode"); mode_layout = QHBoxLayout(mode_group)
        self.mode_conversion_rb = QRadioButton("Convert"); self.mode_conversion_rb.setChecked(True)
        self.mode_encrypt_rb = QRadioButton("Encrypt"); self.mode_decrypt_rb = QRadioButton("Decrypt")
        self.mode_embed_rb = QRadioButton("Embed"); self.mode_extract_rb = QRadioButton("Extract")
        for rb in [self.mode_conversion_rb, self.mode_encrypt_rb, self.mode_decrypt_rb, self.mode_embed_rb, self.mode_extract_rb]:
            mode_layout.addWidget(rb); rb.toggled.connect(self.switch_mode)
        right_layout.addWidget(mode_group)

        # --- Options Groups (stacked for clarity) ---
        self.options_stack = QStackedWidget()
        self.options_stack.addWidget(self._create_conversion_options())
        self.options_stack.addWidget(self._create_encrypt_options())
        self.options_stack.addWidget(self._create_decrypt_options())
        self.options_stack.addWidget(self._create_embed_options())
        self.options_stack.addWidget(self._create_extract_options())
        right_layout.addWidget(self.options_stack)
        
        # --- Output and Run ---
        out_group = QGroupBox("Output"); out_layout = QHBoxLayout(out_group)
        self.outdir_edit = DropLineEdit(self.main_window.settings.get('outdir', DEFAULT_OUTDIR))
        out_layout.addWidget(QLabel("Output folder:")); out_layout.addWidget(self.outdir_edit)
        outbtn = QPushButton("Browse"); outbtn.clicked.connect(lambda: self.outdir_edit.setText(d) if (d:=QFileDialog.getExistingDirectory(self, "Select output directory")) else None)
        out_layout.addWidget(outbtn); right_layout.addWidget(out_group)
        
        self.run_btn = QPushButton("▶️ Run Operation"); self.run_btn.setObjectName("RunButton")
        self.run_btn.clicked.connect(self.start_process)
        right_layout.addWidget(self.run_btn)
        
        right_layout.addStretch()
        splitter.addWidget(right_widget)
        splitter.setSizes([350, 650])
        self.switch_mode() # Set initial UI state

    def _create_conversion_options(self) -> QWidget:
        widget = QGroupBox("Conversion Options"); layout = QVBoxLayout(widget)
        self.conv_type_combo = QComboBox()
        self.conv_type_combo.addItems([
            "PDF (Bank Statements) -> Excel",
            "Word (Tables) -> Excel",
            "Excel -> Combine All Sheets",
            "Combine All Text/CSV Files",
            "Combine All Excel Files",
        ])
        layout.addWidget(self.conv_type_combo)

        self.conv_options_stack = QStackedWidget()
        # PDF to Excel
        pdf_widget = QWidget(); pdf_layout = QFormLayout(pdf_widget)
        self.pdf_filename_edit = QLineEdit("combined_bank_statements")
        self.pdf_flavor_combo = QComboBox(); self.pdf_flavor_combo.addItems(["Stream (Default)", "Lattice"])
        self.pdf_sort_cb = QCheckBox("Sort by date (newest first)"); self.pdf_sort_cb.setChecked(True)
        self.pdf_strict_cb = QCheckBox("Strict Table Extraction"); self.pdf_strict_cb.setChecked(True)
        self.pdf_cols_btn = QPushButton("Preview & Select Columns...")
        self.pdf_cols_btn.clicked.connect(self.open_column_selector)
        pdf_layout.addRow("Output Filename:", self.pdf_filename_edit)
        pdf_layout.addRow("Table Detection Method:", self.pdf_flavor_combo)
        pdf_layout.addRow(self.pdf_sort_cb)
        pdf_layout.addRow(self.pdf_strict_cb)
        pdf_layout.addRow(self.pdf_cols_btn)
        self.conv_options_stack.addWidget(pdf_widget)
        
        # Word to Excel
        word_widget = QWidget(); word_layout = QFormLayout(word_widget)
        word_layout.addRow(QLabel("Each selected .docx file will be converted to a separate .xlsx file."))
        self.word_strict_cb = QCheckBox("Extract Only Tables"); self.word_strict_cb.setChecked(True); self.word_strict_cb.setEnabled(False) # It's the default behavior
        self.word_cols_btn = QPushButton("Preview & Select Columns...")
        self.word_cols_btn.clicked.connect(self.open_column_selector)
        word_layout.addRow(self.word_strict_cb)
        word_layout.addRow(self.word_cols_btn)
        self.conv_options_stack.addWidget(word_widget)

        # Excel Combine Sheets
        excel_sheets_widget = QWidget(); excel_sheets_layout = QVBoxLayout(excel_sheets_widget)
        excel_sheets_layout.addWidget(QLabel("For each selected Excel file, all sheets will be combined into one sheet in a new output file."))
        self.conv_options_stack.addWidget(excel_sheets_widget)

        # Combine Text/CSV
        text_widget = QWidget(); text_layout = QFormLayout(text_widget)
        self.text_filename_edit = QLineEdit("combined_text_files")
        text_layout.addRow("Output Filename:", self.text_filename_edit)
        self.conv_options_stack.addWidget(text_widget)

        # Combine Excel Files
        excel_files_widget = QWidget(); excel_files_layout = QFormLayout(excel_files_widget)
        self.excel_filename_edit = QLineEdit("combined_excel_files")
        self.excel_combine_sheets_cb = QCheckBox("Also combine all sheets within each file"); self.excel_combine_sheets_cb.setChecked(True)
        excel_files_layout.addRow("Output Filename:", self.excel_filename_edit)
        excel_files_layout.addRow(self.excel_combine_sheets_cb)
        self.conv_options_stack.addWidget(excel_files_widget)

        layout.addWidget(self.conv_options_stack)
        self.conv_type_combo.currentIndexChanged.connect(self.conv_options_stack.setCurrentIndex)
        return widget

    def _create_encrypt_options(self) -> QWidget:
        widget = QGroupBox("🌟 Encryptor Options (Fun Mode)"); layout = QFormLayout(widget)
        self.encrypt_ext_edit = QLineEdit(".fun")
        layout.addRow("Disguise with Fake Extension:", self.encrypt_ext_edit)
        return widget

    def _create_decrypt_options(self) -> QWidget:
        widget = QGroupBox("🔓 Decryptor Options (Fun Mode)"); layout = QVBoxLayout(widget)
        layout.addWidget(QLabel("Add encrypted files and run. Original files will be saved to the output folder."))
        return widget

    def _create_embed_options(self) -> QWidget:
        widget = QGroupBox("📦 Embedder Options"); layout = QFormLayout(widget)
        self.host_file_edit = DropLineEdit(); self.host_file_edit.setPlaceholderText("The file to hide data IN (e.g., photo.jpg)")
        self.payload_file_edit = DropLineEdit(); self.payload_file_edit.setPlaceholderText("The file TO HIDE (e.g., secret.txt)")
        layout.addRow("Host File (Drag/Drop):", self.host_file_edit)
        layout.addRow("Payload File (Drag/Drop):", self.payload_file_edit)
        return widget

    def _create_extract_options(self) -> QWidget:
        widget = QGroupBox("🔎 Extractor Options"); layout = QVBoxLayout(widget)
        layout.addWidget(QLabel("Add files with embedded data and run. Payloads will be saved to the output folder."))
        return widget
        
    def open_column_selector(self):
        source_files = self.file_list.get_all_paths()
        if not source_files:
            QMessageBox.warning(self, "No File Selected", "Please add at least one file to the list to use as a preview sample.")
            return

        sample_file = source_files[0]
        conv_type = self.conv_type_combo.currentText()
        file_type = 'pdf' if 'PDF' in conv_type else 'word' if 'Word' in conv_type else None

        if not file_type:
            QMessageBox.information(self, "Not Applicable", "Column selection is only available for PDF and Word to Excel conversions.")
            return

        dialog = TablePreviewDialog(sample_file, file_type, self)
        selected_columns = dialog.get_selected_columns()

        if selected_columns is not None: # User clicked OK
            self.custom_column_config = selected_columns
            QMessageBox.information(self, "Columns Configured", f"Success! This column layout will be applied to the current job:\n\n{', '.join(selected_columns)}")
        else: # User clicked Cancel
            self.custom_column_config = None # Ensure it's cleared if cancelled

    def switch_mode(self):
        if self.mode_conversion_rb.isChecked(): self.options_stack.setCurrentIndex(0)
        elif self.mode_encrypt_rb.isChecked():
            self.options_stack.setCurrentIndex(1)
            if not self.fun_mode_disclaimer_shown:
                QMessageBox.warning(self, "Fun Encryptor", "This is for disguising files, not for security.\nDo not use for sensitive data.")
                self.fun_mode_disclaimer_shown = True
        elif self.mode_decrypt_rb.isChecked(): self.options_stack.setCurrentIndex(2)
        elif self.mode_embed_rb.isChecked(): self.options_stack.setCurrentIndex(3)
        elif self.mode_extract_rb.isChecked(): self.options_stack.setCurrentIndex(4)

    def add_files_dialog(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Select files")
        for f in files: self.file_list.add_file_path(f)
        
    def add_folder_dialog(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder:
            for root, _, files in os.walk(folder):
                for name in files: self.file_list.add_file_path(os.path.join(root, name))
                
    def start_process(self):
        source_files = self.file_list.get_all_paths()
        output_dir = self.outdir_edit.text().strip()

        if not source_files: return QMessageBox.warning(self, "No Files", "Please add files to process.")
        if not output_dir or not os.path.isdir(output_dir):
             os.makedirs(output_dir, exist_ok=True)
             if not os.path.isdir(output_dir):
                return QMessageBox.warning(self, "Invalid Output Directory", "Please specify a valid output directory.")
        
        step: Optional[WorkflowStep] = None
        job_settings = self.main_window.settings.copy()

        if self.mode_conversion_rb.isChecked():
            conv_type = self.conv_type_combo.currentText()
            if conv_type.startswith("PDF"):
                step = PdfToExcelStep()
                step.config['combine_filename'] = self.pdf_filename_edit.text()
                step.config['sort_by_date'] = self.pdf_sort_cb.isChecked()
                step.config['flavor'] = "lattice" if self.pdf_flavor_combo.currentIndex() == 1 else "stream"
            elif conv_type.startswith("Word"):
                step = WordToExcelStep()
            elif conv_type.startswith("Excel -> Combine"):
                step = ExcelProcessStep()
                step.config['combine_sheets'] = True
            elif conv_type.startswith("Combine All Text"):
                step = CombineFilesStep()
                step.config['file_type'] = 'Text'
                step.config['combine_filename'] = self.text_filename_edit.text()
            elif conv_type.startswith("Combine All Excel"):
                step = CombineFilesStep()
                step.config['file_type'] = 'Excel'
                step.config['combine_filename'] = self.excel_filename_edit.text()
                step.config['combine_sheets'] = self.excel_combine_sheets_cb.isChecked()
            
            # Add custom column config if it was set
            if self.custom_column_config:
                step.config['selected_columns'] = self.custom_column_config
                self.custom_column_config = None # Reset for the next job

        elif self.mode_encrypt_rb.isChecked():
            step = FunEncryptStep()
            job_settings['force_ext'] = self.encrypt_ext_edit.text()
        elif self.mode_decrypt_rb.isChecked():
            step = FunDecryptStep()
        elif self.mode_embed_rb.isChecked():
            host = self.host_file_edit.text().strip()
            payload = self.payload_file_edit.text().strip()
            if not host or not payload or not os.path.exists(host) or not os.path.exists(payload):
                return QMessageBox.warning(self, "Error", "Both Host and Payload files must be specified and exist for embedding.")
            source_files = [host]
            step = EmbedFileStep()
            step.config['payload_file'] = payload
        elif self.mode_extract_rb.isChecked():
            step = ExtractFileStep()

        if not step: return QMessageBox.critical(self, "Error", "Could not determine the operation to perform.")

        job_id = f"simple_mode_{int(time.time())}"
        new_job = Job(job_id, source_files, [step], output_dir)
        new_job.workflow_steps[0].config.update(job_settings) # Pass global settings to step

        self.main_window.job_queue.put(new_job)
        self.main_window.log(f"Queued job from Simple Mode: <b>{new_job.get_display_name()}</b>")
        self.main_window.tabs.setCurrentIndex(2) # Switch to the log tab
        QMessageBox.information(self, "Queued", f"The '{step.name}' job has been added to the queue. You can monitor its progress in the 'Job Dashboard' or 'Logs' tab.")

# =====================================================================================
# SECTION 8: MAIN WINDOW & ADVANCED UI
# =====================================================================================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_TITLE); self.resize(1600, 950)
        self.settings_manager = QSettings(SETTINGS_COMPANY, SETTINGS_APP)
        self.db_manager = DatabaseManager(DB_FILE)
        self.job_queue = queue.Queue(); self.failed_jobs: List[Job] = []
        self.worker_thread = None; self.app_worker = None
        self.ai_processor: Optional[AIProcessor] = None
        self.load_settings()
        self.init_ui()
        self.toggle_dark_mode(self.settings.get('dark_mode', True))
        self.db_manager.log_audit_event("App.Startup", "Application started successfully.")

    def get_icon(self, name: str, color: str = '#d3d3d3') -> QIcon:
        return qta.icon(name, color=color) if QTAWESOME_AVAILABLE else QIcon()

    def init_ui(self):
        self.tabs = QTabWidget(); self.setCentralWidget(self.tabs)
        
        # Tab 1: Simple Mode
        self.simple_mode_widget = SimpleModeWidget(self)
        self.tabs.addTab(self.simple_mode_widget, self.get_icon('fa5s.magic'), "Simple Mode")

        # Tab 2: Advanced Workflow Mode
        self.tabs.addTab(self._create_advanced_workflow_tab(), self.get_icon('fa5s.sitemap'), "Advanced Workflow")
        
        # Tab 3: Job Dashboard
        self.tabs.addTab(self._create_job_dashboard_tab(), self.get_icon('fa5s.tasks'), "Job Dashboard")

        # Tab 4: History
        self.tabs.addTab(self._create_history_tab(), self.get_icon('fa5s.history'), "History")
        
        # Tab 5: Logs
        self.tabs.addTab(self._create_log_tab(), self.get_icon('fa5s.file-alt'), "Logs")
        
        self._create_menus()
        self._create_docks()
        self._create_tray_icon()
        self.start_worker_thread()
        self.tabs.setCurrentIndex(0)

    def _create_advanced_workflow_tab(self) -> QWidget:
        main_widget = QWidget(); main_layout = QHBoxLayout(main_widget)
        splitter = QSplitter(Qt.Orientation.Horizontal); main_layout.addWidget(splitter)
        
        # Left Pane: File List & Workflow
        left_pane = QWidget(); left_layout = QVBoxLayout(left_pane)
        file_group = QGroupBox("1. Input Files"); file_layout = QVBoxLayout(file_group)
        self.adv_file_list = FileListWidget(); file_layout.addWidget(self.adv_file_list)
        left_layout.addWidget(file_group)
        
        workflow_group = QGroupBox("2. Build Workflow"); workflow_layout = QVBoxLayout(workflow_group)
        self.workflow_list = QListWidget(); workflow_layout.addWidget(self.workflow_list)
        wf_btn_layout = QHBoxLayout()
        add_step_btn = QPushButton("Add Step"); add_step_btn.clicked.connect(self.add_workflow_step)
        remove_step_btn = QPushButton("Remove Step"); remove_step_btn.clicked.connect(self.remove_workflow_step)
        wf_btn_layout.addWidget(add_step_btn); wf_btn_layout.addWidget(remove_step_btn); workflow_layout.addLayout(wf_btn_layout)
        left_layout.addWidget(workflow_group)
        splitter.addWidget(left_pane)

        # Right Pane: Configuration & Run
        right_pane = QWidget(); right_layout = QVBoxLayout(right_pane)
        self.step_config_group = QGroupBox("3. Configure Step"); self.step_config_layout = QVBoxLayout(self.step_config_group)
        self.step_config_stack = QStackedWidget(); self.step_config_layout.addWidget(self.step_config_stack)
        right_layout.addWidget(self.step_config_group, 1)

        output_group = QGroupBox("4. Output & Execution"); output_layout = QFormLayout(output_group)
        self.adv_outdir_edit = DropLineEdit(self.settings.get('outdir', DEFAULT_OUTDIR))
        output_layout.addRow("Output Directory:", self.adv_outdir_edit)
        self.adv_run_btn = QPushButton("▶️ Queue Workflow Job"); self.adv_run_btn.setObjectName("RunButton")
        self.adv_run_btn.clicked.connect(self.start_advanced_process)
        output_layout.addRow(self.adv_run_btn)
        right_layout.addWidget(output_group)
        splitter.addWidget(right_pane)

        splitter.setSizes([450, 550])
        return main_widget

    def _create_job_dashboard_tab(self) -> QWidget:
        widget = QWidget(); layout = QVBoxLayout(widget)
        queue_group = QGroupBox("Active Job Queue"); queue_layout = QVBoxLayout(queue_group)
        self.job_queue_list = QListWidget(); queue_layout.addWidget(self.job_queue_list)
        layout.addWidget(queue_group)
        failed_group = QGroupBox("Failed Jobs"); failed_layout = QVBoxLayout(failed_group)
        self.failed_jobs_list = QListWidget(); failed_layout.addWidget(self.failed_jobs_list)
        layout.addWidget(failed_group)
        return widget

    def _create_history_tab(self) -> QWidget: # Placeholder
        return QLabel("History Tab (Not shown in this snippet for brevity)")
        
    def _create_log_tab(self) -> QWidget:
        widget = QWidget(); layout = QVBoxLayout(widget)
        self.log_area = QTextEdit(); self.log_area.setReadOnly(True)
        layout.addWidget(self.log_area); return widget

    def _create_menus(self):
        menu_bar = self.menuBar()
        file_menu = menu_bar.addMenu("&File")
        file_menu.addAction("Settings...", self.open_settings)
        file_menu.addSeparator(); file_menu.addAction("Exit", self.close)
        
        view_menu = menu_bar.addMenu("&View")
        self.dark_mode_action = QAction("Dark Mode", self, checkable=True, toggled=self.toggle_dark_mode)
        view_menu.addAction(self.dark_mode_action)
        self.dev_mode_action = QAction("Developer Mode", self, checkable=True, toggled=self.toggle_developer_mode)
        view_menu.addAction(self.dev_mode_action)
        
        tools_menu = menu_bar.addMenu("&Tools")
        tools_menu.addAction("AI Document Query...", self.open_ai_query_dialog)

    def _create_docks(self):
        self.ai_chat_dock = QDockWidget("AI Chat", self)
        self.ai_chat_dock.setWidget(self._create_ai_chat_panel())
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.ai_chat_dock)
        
        self.repl_dock = QDockWidget("Python REPL Console", self)
        if QTCONSOLE_AVAILABLE: self.repl_dock.setWidget(self._create_repl_console())
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.repl_dock)
        self.repl_dock.setVisible(False)

    def _create_ai_chat_panel(self) -> QWidget:
        panel = QWidget(); layout = QVBoxLayout(panel)
        api_layout = QHBoxLayout()
        self.ai_api_type_combo = QComboBox(); self.ai_api_type_combo.addItems(["Google Gemini", "OpenAI", "Anthropic Claude", "Grok (via Groq)"])
        self.ai_api_key_edit = QLineEdit(); self.ai_api_key_edit.setPlaceholderText("Enter API Key..."); self.ai_api_key_edit.setEchoMode(QLineEdit.EchoMode.Password)
        self.ai_init_btn = QPushButton("Initialize AI"); self.ai_init_btn.clicked.connect(self.initialize_ai_processor)
        api_layout.addWidget(self.ai_api_type_combo); api_layout.addWidget(self.ai_api_key_edit, 1); api_layout.addWidget(self.ai_init_btn)
        layout.addLayout(api_layout); return panel
        
    def _create_repl_console(self) -> QWidget:
        kernel_manager = QtInProcessKernelManager(); kernel_manager.start_kernel()
        kernel = kernel_manager.kernel; kernel.shell.push({'main_window': self, 'db': self.db_manager})
        kernel_client = kernel_manager.client(); kernel_client.start_channels()
        jupyter_widget = RichJupyterWidget(); jupyter_widget.kernel_manager = kernel_manager; jupyter_widget.kernel_client = kernel_client
        return jupyter_widget

    def _create_tray_icon(self):
        self.tray_icon = QSystemTrayIcon(self.get_icon('fa5s.cogs'), self)
        tray_menu = QMenu(); tray_menu.addAction("Show", self.show); tray_menu.addAction("Exit", self.close)
        self.tray_icon.setContextMenu(tray_menu); self.tray_icon.show()

    def start_worker_thread(self):
        self.worker_signals = WorkerSignals()
        self.worker_signals.log.connect(self.log)
        self.worker_signals.job_finished.connect(self.on_job_finished)
        self.worker_signals.update_job_dashboard.connect(self.refresh_job_dashboard)
        
        # Initialize AI Processor here but don't configure it yet
        self.ai_processor = AIProcessor(self.worker_signals)
        
        self.app_worker = AppWorker(self.job_queue, self.worker_signals, self.db_manager, self.ai_processor)
        self.worker_thread = QThread(); self.app_worker.moveToThread(self.worker_thread)
        self.worker_thread.started.connect(self.app_worker.run); self.worker_thread.start()
        self.log("Worker thread started and ready for jobs.")

    # --- Event Handlers & Logic ---
    def open_settings(self):
        dialog = SettingsDialog(self.settings, self)
        if dialog.exec(): self.settings = dialog.get_settings()
        
    def open_ai_query_dialog(self):
        if not self.ai_processor or not self.ai_processor.is_ready():
            QMessageBox.warning(self, "AI Not Ready", "Please initialize an AI provider in the AI Chat dock first.")
            return
        dialog = AIQueryDialog(self.ai_processor, self)
        dialog.exec()

    def add_workflow_step(self):
        step_name, ok = QInputDialog.getItem(self, "Add Workflow Step", "Select step type:", ALL_WORKFLOW_STEPS.keys(), 0, False)
        if ok and step_name:
            step_instance = ALL_WORKFLOW_STEPS[step_name]()
            config_widget = step_instance.get_config_widget()
            self.step_config_stack.addWidget(config_widget)
            item = QListWidgetItem(step_name)
            item.setData(Qt.ItemDataRole.UserRole, step_instance)
            self.workflow_list.addItem(item)
            
    def remove_workflow_step(self):
        if (row := self.workflow_list.currentRow()) > -1:
            self.workflow_list.takeItem(row) # Simplified removal
            
    def start_advanced_process(self):
        source_files = self.adv_file_list.get_all_paths()
        workflow_steps = [self.workflow_list.item(i).data(Qt.ItemDataRole.UserRole) for i in range(self.workflow_list.count())]
        if not source_files or not workflow_steps:
                 return QMessageBox.warning(self, "Missing Info", "Please add files and at least one workflow step.")
        job_id = f"adv_workflow_{int(time.time())}"
        new_job = Job(job_id, source_files, workflow_steps, self.adv_outdir_edit.text())
        self.job_queue.put(new_job)
        self.log(f"Queued advanced workflow job: <b>{new_job.get_display_name()}</b>")
        self.tabs.setCurrentIndex(2) # Switch to dashboard

    def log(self, message: str):
        if hasattr(self, 'log_area'):
            QMetaObject.invokeMethod(self.log_area, "append", Qt.ConnectionType.QueuedConnection, Q_ARG(str, message))
    
    def on_job_finished(self, job: Job):
        if job.status == "failed": self.failed_jobs.append(job)
        self.refresh_job_dashboard()
        self.tray_icon.showMessage(f"Job {job.status.title()}", f"'{job.get_display_name()}' has finished.")
        
    def refresh_job_dashboard(self):
        self.job_queue_list.clear()
        with self.job_queue.mutex:
            for job in list(self.job_queue.queue): self.job_queue_list.addItem(f"{job.get_display_name()} [{job.status}]")
        self.failed_jobs_list.clear()
        for job in self.failed_jobs: self.failed_jobs_list.addItem(f"{job.get_display_name()} - Error: {job.error_message}")
    
    def initialize_ai_processor(self):
        api_type = self.ai_api_type_combo.currentText(); api_key = self.ai_api_key_edit.text().strip()
        if not api_key: return QMessageBox.warning(self, "API Key Missing", "Please enter a valid API key.")
        
        if self.ai_processor.initialize(api_type, api_key):
            self.ai_init_btn.setEnabled(False); self.ai_api_key_edit.setEnabled(False)
            self.ai_api_type_combo.setEnabled(False)
            QMessageBox.information(self, "AI Initialized", f"Successfully connected to {api_type}.")
        else:
            QMessageBox.critical(self, "AI Initialization Failed", "Could not connect to the AI service. Please check your API key and network connection. See logs for details.")

    def load_settings(self):
        self.settings = {
            'outdir': self.settings_manager.value('outdir', DEFAULT_OUTDIR, type=str),
            'dark_mode': self.settings_manager.value('dark_mode', True, type=bool),
            'gpu_device': self.settings_manager.value('gpu_device', 'auto', type=str),
            'chunk_size_mb': self.settings_manager.value('chunk_size_mb', 64, type=int),
        }
    
    def save_settings(self):
        self.settings_manager.setValue('outdir', self.simple_mode_widget.outdir_edit.text())
        self.settings_manager.setValue('dark_mode', self.dark_mode_action.isChecked())
        for key in ['gpu_device', 'chunk_size_mb']:
            if key in self.settings: self.settings_manager.setValue(key, self.settings[key])

    def toggle_dark_mode(self, checked):
        # Stylesheet for a modern dark theme
        dark_stylesheet = """
            QWidget { background-color: #2b2b2b; color: #f0f0f0; border: none; }
            QMainWindow, QDockWidget { border: 1px solid #555; }
            QGroupBox { border: 1px solid #444; margin-top: 1ex; padding-top: 5px; }
            QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }
            QLineEdit, QTextEdit, QSpinBox, QTableWidget { background-color: #3c3f41; border: 1px solid #555; border-radius: 4px; padding: 3px; }
            QPushButton { background-color: #4a4d4f; border: 1px solid #666; padding: 5px; border-radius: 4px; }
            QPushButton:hover { background-color: #5a5d5f; }
            QPushButton:pressed { background-color: #3a3d3f; }
            QPushButton#RunButton { font-size: 16px; padding: 10px; background-color: #0d6efd; color: white; }
            QListWidget, QComboBox { background-color: #3c3f41; border: 1px solid #555; }
            QListWidget::item:selected { background-color: #0d6eff; }
            QTabWidget::pane { border-top: 1px solid #444; }
            QTabBar::tab { padding: 8px 20px; background: #2b2b2b; border: 1px solid #2b2b2b; border-bottom: none; }
            QTabBar::tab:hover { background: #3c3f41; }
            QTabBar::tab:selected { background: #3c3f41; border: 1px solid #444; border-bottom: 1px solid #3c3f41; }
            QHeaderView::section { background-color: #4a4d4f; padding: 4px; border: 1px solid #555; }
        """
        QApplication.instance().setStyleSheet(dark_stylesheet if checked else "")
        
    def toggle_developer_mode(self, checked):
        self.repl_dock.setVisible(checked)

    def closeEvent(self, event):
        self.save_settings()
        if self.app_worker: self.app_worker.stop()
        if self.worker_thread: self.worker_thread.quit(); self.worker_thread.wait(3000)
        event.accept()

# =====================================================================================
# SECTION 9: INTERACTIVE DIALOGS (COLUMN SELECTOR)
# =====================================================================================

class TablePreviewDialog(QDialog):
    def __init__(self, file_path: str, file_type: str, parent=None):
        super().__init__(parent)
        self.file_path = file_path
        self.file_type = file_type
        self.preview_df = None
        self.setWindowTitle(f"Column Selector - {os.path.basename(file_path)}")
        self.setMinimumSize(800, 600)

        self.init_ui()
        self.load_preview_data()

    def init_ui(self):
        main_layout = QVBoxLayout(self)
        
        self.table_preview = QTableWidget()
        self.table_preview.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        main_layout.addWidget(QLabel("Table Preview (First table found in document)"), 0)
        main_layout.addWidget(self.table_preview, 1) # Give more space to table

        selector_layout = QHBoxLayout()
        # Available columns
        available_group = QGroupBox("Available Columns")
        available_layout = QVBoxLayout(available_group)
        self.available_list = QListWidget()
        self.available_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        available_layout.addWidget(self.available_list)
        selector_layout.addWidget(available_group)

        # Control buttons
        controls_layout = QVBoxLayout()
        controls_layout.addStretch()
        add_btn = QPushButton(">>")
        add_btn.setToolTip("Add selected columns")
        add_btn.clicked.connect(self.add_columns)
        remove_btn = QPushButton("<<")
        remove_btn.setToolTip("Remove selected columns")
        remove_btn.clicked.connect(self.remove_columns)
        controls_layout.addWidget(add_btn)
        controls_layout.addWidget(remove_btn)
        controls_layout.addStretch()
        selector_layout.addLayout(controls_layout)

        # Selected columns
        selected_group = QGroupBox("Selected Columns (in order)")
        selected_layout = QVBoxLayout(selected_group)
        self.selected_list = QListWidget()
        self.selected_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        reorder_btns = QHBoxLayout()
        up_btn = QPushButton("Up")
        up_btn.clicked.connect(self.move_up)
        down_btn = QPushButton("Down")
        down_btn.clicked.connect(self.move_down)
        reorder_btns.addStretch(); reorder_btns.addWidget(up_btn); reorder_btns.addWidget(down_btn)
        selected_layout.addWidget(self.selected_list)
        selected_layout.addLayout(reorder_btns)
        selector_layout.addWidget(selected_group)
        
        main_layout.addLayout(selector_layout, 1) # Give space to selector too

        # OK/Cancel buttons
        button_box = QHBoxLayout()
        ok_button = QPushButton("OK"); ok_button.clicked.connect(self.accept)
        cancel_button = QPushButton("Cancel"); cancel_button.clicked.connect(self.reject)
        button_box.addStretch(); button_box.addWidget(ok_button); button_box.addWidget(cancel_button)
        main_layout.addLayout(button_box)

    def load_preview_data(self):
        try:
            if self.file_type == 'pdf':
                tables = camelot.read_pdf(self.file_path, pages='1', flavor='stream')
                if tables.n > 0: self.preview_df = tables[0].df
            elif self.file_type == 'word':
                doc = Document(self.file_path)
                if doc.tables:
                    table = doc.tables[0]
                    data = [[cell.text for cell in row.cells] for row in table.rows]
                    if len(data) > 1:
                        self.preview_df = pd.DataFrame(data[1:], columns=data[0])
                    else: # Handle empty or header-only table
                        self.preview_df = pd.DataFrame(columns=data[0] if data else [])
            
            if self.preview_df is not None:
                self.populate_widgets()
            else:
                QMessageBox.warning(self, "No Tables Found", "Could not find any tables in the first page/section of the sample document.")
        except Exception as e:
            QMessageBox.critical(self, "Error Loading Preview", f"An error occurred: {e}")

    def populate_widgets(self):
        # Populate QTableWidget
        self.preview_df = self.preview_df.head(20) # Limit to first 20 rows for preview
        self.table_preview.setRowCount(self.preview_df.shape[0])
        self.table_preview.setColumnCount(self.preview_df.shape[1])
        self.table_preview.setHorizontalHeaderLabels(self.preview_df.columns.astype(str))
        for i, row in self.preview_df.iterrows():
            for j, val in enumerate(row):
                self.table_preview.setItem(i, j, QTableWidgetItem(str(val)))
        self.table_preview.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)

        # Populate Available Columns list
        self.available_list.addItems(self.preview_df.columns.astype(str))

    def add_columns(self):
        selected_items = self.available_list.selectedItems()
        for item in selected_items:
            # Check if item already exists in selected list to avoid duplicates
            if not self.selected_list.findItems(item.text(), Qt.MatchFlag.MatchExactly):
                self.selected_list.addItem(item.text())

    def remove_columns(self):
        for item in self.selected_list.selectedItems():
            self.selected_list.takeItem(self.selected_list.row(item))

    def move_up(self):
        row = self.selected_list.currentRow()
        if row > 0:
            item = self.selected_list.takeItem(row)
            self.selected_list.insertItem(row - 1, item)
            self.selected_list.setCurrentRow(row - 1)

    def move_down(self):
        row = self.selected_list.currentRow()
        if row < self.selected_list.count() - 1:
            item = self.selected_list.takeItem(row)
            self.selected_list.insertItem(row + 1, item)
            self.selected_list.setCurrentRow(row + 1)

    def get_selected_columns(self) -> Optional[List[str]]:
        if self.exec() == QDialog.DialogCode.Accepted:
            return [self.selected_list.item(i).text() for i in range(self.selected_list.count())]
        return None

# =====================================================================================
# SECTION 10: CLI & EXECUTION START
# =====================================================================================

def main():
    app = QApplication(sys.argv)
    app.setOrganizationName(SETTINGS_COMPANY); app.setApplicationName(SETTINGS_APP)
    main_window = MainWindow()
    main_window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    if sys.platform == "win32":
        import ctypes
        myappid = f'{SETTINGS_COMPANY}.{SETTINGS_APP}.2.3' # Updated version
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    main()